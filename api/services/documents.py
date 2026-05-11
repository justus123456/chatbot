from io import BytesIO

import fitz
from docx import Document as DocxDocument


def extract_text(file_storage):
    filename = (file_storage.filename or "").lower()
    data = file_storage.read()
    file_storage.seek(0)
    if filename.endswith(".pdf"):
        with fitz.open(stream=data, filetype="pdf") as document:
            return "\n".join(page.get_text() for page in document)
    if filename.endswith(".docx"):
        document = DocxDocument(BytesIO(data))
        parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts)
    return data.decode("utf-8", errors="ignore")


def chunk_text(text, size=1800, overlap=200):
    clean = " ".join(text.split())
    chunks = []
    start = 0
    while start < len(clean):
        end = min(start + size, len(clean))
        chunks.append(clean[start:end])
        if end == len(clean):
            break
        start = max(0, end - overlap)
    return chunks
